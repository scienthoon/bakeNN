#!/usr/bin/env python3
"""Build the standalone BakeNN contest result-report draft."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs/submission/BakeNN_오픈소스개발자대회_결과보고서_초안.docx"
FONT = "Arial Unicode MS"  # named Korean-font override for standard_business_brief
MONO = "Menlo"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "666666"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 80, bottom: int = 80,
                      start: int = 120, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("bottom", bottom),
                        ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: tuple[int, ...]) -> None:
    if sum(widths) != TABLE_WIDTH:
        raise ValueError(f"table widths must total {TABLE_WIDTH}: {widths}")
    table.autofit = False
    properties = table._tbl.tblPr
    width_node = properties.find(qn("w:tblW"))
    if width_node is None:
        width_node = OxmlElement("w:tblW")
        properties.append(width_node)
    width_node.set(qn("w:w"), str(TABLE_WIDTH))
    width_node.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, value in zip(row.cells, widths):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(value))
            tc_width.set(qn("w:type"), "dxa")


def _set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _font_run(run, size: float = 11, color: str = "000000",
              bold: bool = False, italic: bool = False,
              font: str = FONT) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _style_font(style, size: float, color: str = "000000",
                bold: bool = False) -> None:
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 30, INK, 0, 8),
        ("Subtitle", 14, GRAY, 0, 12),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        _style_font(style, size, color, name != "Subtitle")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    _style_font(caption, 9, GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def _numbering(document: Document, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId")))
                    for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId")))
               for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def _add_list_item(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_properties.append(level)
    num_properties.append(number)
    properties.append(num_properties)
    _font_run(paragraph.add_run(text))


def _add_table(document: Document, headers: tuple[str, ...],
               rows: tuple[tuple[str, ...], ...], widths: tuple[int, ...],
               numeric_columns: tuple[int, ...] = ()):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_geometry(table, widths)
    _set_repeat_header(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                               if index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.paragraph_format.space_after = Pt(0)
        _font_run(paragraph.add_run(value), size=9.5, color=INK, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                                   if index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT)
            paragraph.paragraph_format.space_after = Pt(0)
            _font_run(paragraph.add_run(value), size=9.3)
    _set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _add_callout(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_BLUE)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    properties.append(borders)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    _font_run(paragraph.add_run(f"{label}  "), size=10.5, color=INK, bold=True)
    _font_run(paragraph.add_run(text), size=10.5, color=INK)


def _field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, end))
    _font_run(run, size=9, color=GRAY)


def _configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    _font_run(header.add_run("BakeNN · 오픈소스 개발자대회 결과보고서"),
              size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font_run(footer.add_run("BakeNN  |  "), size=9, color=GRAY)
    _field(footer, "PAGE")
    _font_run(footer.add_run(" / "), size=9, color=GRAY)
    _field(footer, "NUMPAGES")


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def _add_paragraph(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        _font_run(paragraph.add_run(bold_lead), bold=True)
        _font_run(paragraph.add_run(text[len(bold_lead):]))
    else:
        _font_run(paragraph.add_run(text))


def build() -> Path:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    bullets = _numbering(document, bullet=True)
    numbers = _numbering(document, bullet=False)

    # editorial_cover header pattern, standard_business_brief tokens
    for _ in range(4):
        document.add_paragraph().paragraph_format.space_after = Pt(12)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font_run(kicker.add_run("오픈소스 개발자대회 · 결과보고서 초안"),
              size=11, color=BLUE, bold=True)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font_run(title.add_run("BakeNN"), size=30, color=INK, bold=True)
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font_run(subtitle.add_run("고정 모델 MCU를 위한 INT8 AOT 컴파일러"),
              size=15, color=DARK_BLUE)
    lead = document.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font_run(
        lead.add_run(
            "학습된 PyTorch FP32 모델과 대표 데이터를 받아\n"
            "인터프리터 없는 정적 C11 추론 라이브러리를 생성"
        ),
        size=11,
        color=GRAY,
    )
    document.add_paragraph().paragraph_format.space_after = Pt(22)
    _add_table(
        document,
        ("항목", "내용"),
        (
            ("저장소", "https://github.com/scienthoon/bakeNN"),
            ("라이선스", "Apache License 2.0"),
            ("버전", "BakeNN 0.1.0 (출품 준비 브랜치)"),
            ("작성 기준일", "2026년 8월 16일"),
            ("작성자", "scienthoon (GitHub)"),
        ),
        (2700, 6660),
    )
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font_run(
        note.add_run("공식 결과보고서 양식 수령 시 본문의 항목을 해당 서식에 이관"),
        size=9.5,
        color=GRAY,
        italic=True,
    )
    document.add_page_break()

    _add_heading(document, "1. 출품작 개요")
    _add_callout(
        document,
        "핵심 결과",
        "BakeNN은 정적 batch-1 PyTorch FP32 모델을 대표 데이터로 PTQ하고, "
        "검증된 INT8 실행 계획·SRAM 주소·커널 선택을 확정한 뒤 모델 전용 C11을 생성한다.",
    )
    _add_paragraph(
        document,
        "BakeNN은 MCU 펌웨어에 한두 개의 고정 모델을 포함하고 모델 변경 시 펌웨어도 "
        "함께 갱신하는 양산 제품을 대상으로 한다. 런타임에서 FlatBuffer를 해석하지 않고 "
        "호스트 컴파일 단계에서 그래프, 텐서 수명, 양자화 파라미터, 메모리 주소와 커널을 "
        "결정한다. 결과물은 model.h/model.c/weights/manifest와 필요한 커널 소스뿐이다.",
    )
    _add_paragraph(
        document,
        "정적 batch 1과 단일 공개 입력·출력 ABI는 의도적인 제품 제약이다. 내부 그래프는 "
        "residual, fan-out, concatenate, broadcast와 skip connection을 포함할 수 있지만, "
        "동적 모델 로딩과 임의 shape를 버림으로써 컴파일 시점 메모리 보증과 직접 호출 코드를 얻는다.",
    )

    _add_heading(document, "2. 개발 배경과 해결하려는 문제")
    for item in (
        "TFLite Micro는 넓은 연산 범위와 런타임 모델 교체 능력을 제공하지만, 고정 제품에서도 FlatBuffer, MicroInterpreter, op resolver, tensor planner와 런타임 메타데이터를 유지한다.",
        "모델이 바뀌면 resolver 용량·연산 등록, operator version, TFLM 및 vendor kernel 버전의 호환성을 개발자가 함께 관리해야 한다.",
        "모델 SRAM이 부족한 문제는 보드에서 AllocateTensors 실패로 늦게 발견될 수 있고, 최종 실행 순서를 감사하려면 모델 바이너리와 런타임 여러 계층을 함께 추적해야 한다.",
        "BakeNN은 범용 런타임 대신 고정 그래프를 강하게 정적 컴파일하여 이 조립·해석 비용을 호스트 단계로 이동한다.",
    ):
        _add_list_item(document, item, bullets)

    _add_heading(document, "3. 설계와 컴파일 파이프라인")
    architecture = document.add_paragraph()
    architecture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    architecture.paragraph_format.space_before = Pt(8)
    architecture.paragraph_format.space_after = Pt(10)
    _font_run(
        architecture.add_run(
            "PyTorch FP32 eval 모델 + calibration samples\n"
            "↓  torch.export 캡처 · BatchNorm folding\n"
            "FloatGraph → PTQ QuantizedGraph → legalize/fusion\n"
            "↓  검증 · liveness · static SRAM allocation\n"
            "ExecutionPlan → backend kernel selection\n"
            "↓\nmodel.h + model.c + weights + manifest + 필요한 커널"
        ),
        size=10,
        color=INK,
        font=MONO,
    )
    for item in (
        "프런트엔드: torch.export의 정적 의미를 framework-neutral FloatGraph로 정규화하며 eval BatchNorm과 Dropout 등을 컴파일 시 제거·결합한다.",
        "PTQ: 대표 입력으로 각 중간 activation 범위를 관측하고 per-tensor affine activation, per-output-channel symmetric weight, int32 bias를 결정한다.",
        "정수 IR: shape·layout·dtype·scale·zero point를 edge별로 고정하고 지원하지 않는 의미는 CompileError로 거절한다.",
        "정적 계획: 텐서 생존 구간을 계산해 activation buffer를 재사용하고 vendor scratch까지 하나의 정적 arena에 배치한다.",
        "코드 생성: 실행 순서와 주소가 상수인 C11 함수와 모델에 필요한 커널 소스만 방출한다.",
    ):
        _add_list_item(document, item, numbers)

    _add_heading(document, "4. INT8 수치 계약과 정확성")
    _add_table(
        document,
        ("항목", "BakeNN 0.1.0 계약"),
        (
            ("Activation", "per-tensor affine int8 [-128, 127]"),
            ("Weight", "output-channel별 symmetric int8 [-127, 127], zero point 0"),
            ("Bias/Accumulator", "int32, host에서 accumulator overflow 상한 증명"),
            ("Requantization", "versioned Q31 double-round profile: bakenn.int8.v1"),
            ("Padding", "실수 0에 대응하는 input zero point 사용"),
            ("검증 기준", "Python integer reference ↔ portable/vendor C output byte-exact"),
        ),
        (2700, 6660),
    )
    _add_paragraph(
        document,
        "변환 후 Python 정수 실행기와 생성 C는 같은 고정소수점 계약을 사용한다. "
        "Q31 multiplier·shift는 호스트에서 한 번 계산해 상수로 내보내며 타겟에서 float로 재계산하지 않는다. "
        "지원 범위를 벗어난 vendor kernel은 AUTO에서 portable로 fallback하고 REQUIRE_OPTIMIZED에서는 이유와 함께 실패한다.",
    )

    _add_heading(document, "5. 구현 범위")
    _add_table(
        document,
        ("분류", "지원 내용", "대표 모델/용도"),
        (
            ("핵심 연산", "Conv2D/Depthwise/Linear, grouped Conv, Conv1D, grouped ConvTranspose2D", "CNN, audio/time-series, 정적 U-Net"),
            ("Activation", "ReLU/Clamp, Sigmoid, HardSigmoid, HardSwish, SiLU", "MobileNetV3, EfficientNet 계열"),
            ("Pool/Reduction", "Avg/Max 1D·2D, ReduceMean, global pooling", "분류 head, temporal model"),
            ("Tensor/Shape", "Reshape/Flatten/Squeeze, static Slice/Crop, Concatenate, Pad, Resize", "Residual, Dense, encoder-decoder"),
            ("Elementwise", "Add/Mul과 정적-rank broadcast, internal Requantize", "SE block, residual connection"),
            ("출력", "BakeNN Q15-LUT Softmax", "고정 class classifier"),
        ),
        (1800, 4860, 2700),
    )
    _add_paragraph(
        document,
        "호스트 모델 게이트는 MobileNetV2/V3, EfficientNet-Lite-style, ResNet bottleneck, "
        "DenseNet-style concat, Inception, SqueezeNet Fire, compact U-Net, Conv1D classifier와 "
        "temporal residual 모델을 포함한다. TFLM 전체 연산 카탈로그와 동등하다는 의미는 아니다.",
    )

    _add_heading(document, "6. 타겟과 최적화 백엔드")
    _add_table(
        document,
        ("타겟", "실행 경로", "현재 검증"),
        (
            ("범용 32-bit MCU", "heap-free portable C11", "GCC/Clang + ASan/UBSan"),
            ("Cortex-M0+", "portable C fallback", "arm-none-eabi cross-link"),
            ("Cortex-M4 DSP", "SMLAD kernels + direct CMSIS-NN FC/Conv/DW/Pool", "cross-ELF/disassembly + nRF52840 실측"),
            ("RV32IMC", "portable/generic C fallback", "riscv GNU cross-link"),
            ("ESP32", "ESP-NN optimized Conv/DW + fallback", "official C differential + ESP-IDF build"),
            ("ESP32-S3", "ESP-NN SIMD Conv/DW/FC/Pool + fallback", "ANSI oracle differential + Xtensa ESP-IDF build"),
            ("ESP32-C3", "portable/generic fallback", "ESP-IDF build"),
        ),
        (1900, 4100, 3360),
    )

    _add_heading(document, "7. 검증 결과")
    _add_callout(
        document,
        "전체 회귀",
        "2026-08-16 동결 브랜치에서 279 tests + 6 subtests가 통과했다. "
        "GitHub Actions는 Python 3.10/3.12 × GCC/Clang, PyTorch frontend/wheel, ARM/RISC-V cross-toolchain, ESP-IDF 3종을 검증한다.",
    )
    for item in (
        "hand-calculated golden, malformed IR, qparam/overflow/alias 검증과 randomized differential를 함께 수행한다.",
        "대표 생성 C는 strict C11, -Wall -Wextra -Werror -pedantic, ASan, UBSan으로 컴파일·실행한다.",
        "ESP32 optimized Conv/Depthwise는 10,000 random/edge tensor에서 Python reference와 byte mismatch 0이었다.",
        "ESP32-S3 Conv/Depthwise/FC/Pool wrapper는 graph별 10,000 tensor를 official ANSI oracle 경유로 비교해 mismatch 0이었다.",
        "wheel에는 pinned CMSIS-NN/ESP-NN source와 license/provenance가 포함되며 fresh venv 설치 후 모델 컴파일을 재검증했다.",
    ):
        _add_list_item(document, item, bullets)

    _add_heading(document, "8. 실제 학습 모델 PTQ 결과")
    _add_paragraph(
        document,
        "MNIST 2종과 CIFAR-10 4종을 각각 전체 training split으로 1 epoch 학습하고, "
        "100개 class-balanced 이미지로 calibration한 뒤 10,000개 test 이미지를 생성 C로 실행했다.",
    )
    _add_table(
        document,
        ("모델", "Dataset", "FP32", "INT8 C", "FP32-INT8"),
        (
            ("MNIST MLP", "MNIST", "92.69%", "92.68%", "+0.01 pp"),
            ("MNIST CNN", "MNIST", "92.53%", "92.52%", "+0.01 pp"),
            ("Bottleneck", "CIFAR-10", "20.20%", "20.22%", "-0.02 pp"),
            ("Dense concat", "CIFAR-10", "18.83%", "19.34%", "-0.51 pp"),
            ("Inception", "CIFAR-10", "18.64%", "18.36%", "+0.28 pp"),
            ("Fire", "CIFAR-10", "18.59%", "18.69%", "-0.10 pp"),
        ),
        (2200, 1800, 1650, 1650, 2060),
        numeric_columns=(2, 3, 4),
    )
    _add_paragraph(
        document,
        "여섯 실행 모두 Python INT8과 생성 C 비교에서 output byte mismatch가 0이었다. "
        "CIFAR-10 절대 정확도가 낮은 이유는 의도적으로 작은 모델을 단 1 epoch만 학습했기 때문이며, "
        "본 실험은 SOTA 정확도가 아니라 실제 학습→calibration→PTQ→C 파이프라인과 양자화 차이를 검증한다.",
    )

    _add_heading(document, "9. TFLite Micro 대비 실보드 결과")
    _add_paragraph(
        document,
        "동일한 32→16→4 INT8 FC workload를 IoT-LAB nRF52840DK 64 MHz에서 "
        "동일 qparams·weights·bias·input bytes·output 의미로 비교했다. 8 warmups 후 101회를 측정했다.",
    )
    _add_table(
        document,
        ("Build", "Median cycles", "Flash", "Static SRAM"),
        (
            ("BakeNN direct CMSIS-NN", "3,786", "20,920 B", "8,540 B"),
            ("TFLM + CMSIS-NN", "5,418", "69,640 B", "11,040 B"),
            ("BakeNN portable", "8,706", "20,764 B", "8,540 B"),
            ("TFLM reference", "9,342", "63,176 B", "11,008 B"),
        ),
        (3900, 1800, 1800, 1860),
        numeric_columns=(1, 2, 3),
    )
    _add_callout(
        document,
        "동일 FC workload 결과",
        "BakeNN direct CMSIS-NN은 TFLM+CMSIS-NN 대비 cycles 30.1% 감소, "
        "linked Flash 70.0% 감소, linked static SRAM 22.6% 감소했다. 네 이미지의 output bytes와 FNV-1a 0x910c1fe2가 일치했다.",
    )
    _add_paragraph(
        document,
        "별도 1×4×4×1→1×4×4×2 Conv2D에서 BakeNN portable은 TFLM reference 대비 "
        "24,610 대 27,441 cycles, 20,332 대 61,760 B Flash, 8,160 대 10,624 B static SRAM을 기록했다. "
        "이는 CMSIS-NN Conv 비교가 아니며 모든 모델·MCU로 일반화하지 않는다.",
    )

    _add_heading(document, "10. TFLM 대비 실용적 차별점")
    differentiator_numbers = _numbering(document, bullet=False)
    for item in (
        "외부 model interpreter를 MCU에 이식하지 않고 생성 C와 선택된 kernel만 MCU compiler로 빌드한다.",
        "shape·padding·qparams·주소·실행 순서가 상수이므로 fusion, buffer reuse, packing과 전용 kernel 선택을 모델별로 수행한다.",
        "constant, activation arena, scratch와 alignment를 호스트에서 계산해 Flash/SRAM budget 초과를 보드에 올리기 전 CI에서 거절한다.",
        "생성 C 호출 순서와 manifest를 직접 검토할 수 있어 FlatBuffer·resolver·planner·runtime을 함께 추적하는 것보다 펌웨어 감사 범위가 작다.",
        "실제 비교 과정에서 TFLM은 resolver 등록, Conv2D operator version, arena reservation, 구버전 CMSIS-NN wrapper 호환을 수동으로 맞춰야 했지만 BakeNN은 검증된 graph에서 kernel/source closure를 자동 확정했다.",
    ):
        _add_list_item(document, item, differentiator_numbers)
    _add_paragraph(
        document,
        "대가로 BakeNN은 TFLM보다 연산 범위가 좁고, 모델을 바꾸려면 펌웨어를 다시 빌드해야 하며, "
        "동적 shape·다중 공개 입출력·런타임 모델 교체를 지원하지 않는다. 이러한 요구가 중요한 제품에는 TFLM이 더 적합하다.",
    )

    _add_heading(document, "11. 오픈소스 구성과 재현 방법")
    _add_table(
        document,
        ("항목", "내용"),
        (
            ("Source", "https://github.com/scienthoon/bakeNN"),
            ("License", "Apache-2.0; 상업 이용·수정·재배포 허용"),
            ("CI", "GCC/Clang, PyTorch, wheel, ARM/RISC-V, ESP-IDF"),
            ("Community", "CONTRIBUTING, CODE_OF_CONDUCT, SECURITY, issue/PR templates"),
            ("Evidence", "benchmarks/RESULTS.md 및 원본 UART/JSON report"),
            ("Demo", "examples/esp32s3_end_to_end/"),
        ),
        (2300, 7060),
    )
    reproduction = document.add_paragraph()
    reproduction.paragraph_format.left_indent = Inches(0.25)
    reproduction.paragraph_format.space_before = Pt(6)
    reproduction.paragraph_format.space_after = Pt(8)
    _font_run(
        reproduction.add_run(
            "git clone https://github.com/scienthoon/bakeNN.git\n"
            "cd bakeNN\n"
            "python -m pip install -e '.[test,torch]'\n"
            "pytest -q\n"
            "python examples/esp32s3_end_to_end/generate.py --output build/esp32s3_demo"
        ),
        size=9.5,
        color=INK,
        font=MONO,
    )

    _add_heading(document, "12. 한계와 향후 계획")
    _add_table(
        document,
        ("현재 한계", "후속 작업"),
        (
            ("ESP32/S3 물리 cycle·energy 미측정", "동일 모델 TFLM/ESP-NN 실보드 비교와 measured AUTO cost table"),
            ("TFLM보다 좁은 operator surface", "시장별 fixture 기반으로 detection/postprocess 등 선택 확장"),
            ("PTQ만 구현, QAT fine-tuning 없음", "PTQ 정확도가 부족한 모델을 위한 별도 QAT frontend"),
            ("동적 shape·다중 공개 I/O 없음", "핵심 제품 계약은 유지하고 필요한 정적 multi-ABI를 별도 검토"),
            ("full application stack peak 미증명", "최종 firmware ELF, stack watermark, RTOS/global 포함 측정"),
        ),
        (3800, 5560),
    )

    _add_heading(document, "13. 제출 항목 상태")
    _add_table(
        document,
        ("필수 항목", "상태", "비고"),
        (
            ("소스코드", "준비", "public GitHub, Apache-2.0"),
            ("결과보고서 원본", "준비", "본 DOCX 초안; 공식 양식에 이관 필요"),
            ("결과보고서 PDF", "준비", "본 문서 렌더 PDF"),
            ("시연 절차", "준비", "3분 demo script와 ESP32-S3 one-command generator"),
            ("시연영상 링크", "미완료", "화면 녹화·업로드 후 공식 양식과 본 문서에 링크 추가 필요"),
            ("최종 release/tag", "미완료", "green PR merge 후 main에서 v0.1.0 tag/release"),
        ),
        (2600, 1500, 5260),
        numeric_columns=(1,),
    )
    _add_callout(
        document,
        "제출 전 필수",
        "시연영상 링크를 추가하고 공식 주최측 양식에 본문을 이관해야 한다. "
        "마감 전에 홈페이지 상태가 ‘제출 완료’로 바뀌고 자동 안내 메일이 도착했는지 모두 확인한다.",
    )

    document.add_page_break()
    _add_heading(document, "부록 A. 증거 파일")
    for item in (
        "benchmarks/tflm_compare/results/iotlab_447626_direct_cmsis_fc.md",
        "benchmarks/tflm_compare/results/iotlab_447626_direct_cmsis_fc_uart.txt",
        "benchmarks/tflm_compare/results/iotlab_447609_conv.json",
        "examples/training_matrix/RESULTS.md",
        "benchmarks/RESULTS.md",
        "docs/P0_STATUS.md",
        "docs/P2_KERNEL_ARCHITECTURE.md",
        "docs/RELEASE_CHECKLIST.md",
    ):
        _add_list_item(document, item, bullets)

    document.core_properties.title = "BakeNN 오픈소스 개발자대회 결과보고서 초안"
    document.core_properties.subject = "고정 모델 MCU용 INT8 AOT 컴파일러"
    document.core_properties.author = "scienthoon"
    document.core_properties.keywords = "BakeNN, INT8, AOT, MCU, PyTorch, TFLM"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
